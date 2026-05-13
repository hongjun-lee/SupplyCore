"""
批量后处理发标包 docx 样式：
  1. 正文：仿宋_GB2312 小三 (15pt)
  2. 文章标题（Title / Heading 1）：仿宋 三号 (16pt) 加粗 居中
  3. 章节标题（Heading 2~6）：仿宋 小三 (15pt) 加粗
  4. 表格单元格四周加实体边框（single line, 4pt 宽度, 黑色），单元格文字同正文样式

用法:
  python3 scripts/apply_docx_style.py "docs/招标/word"
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


ZH_FONT_BODY = "仿宋_GB2312"
ZH_FONT_HEADING = "仿宋"
EN_FONT = "Times New Roman"
BASE_PT = 15  # 小三

# 标题样式配置：字号 / 中文字体 / 是否居中（统一加粗）
HEADING_CONFIG: dict[str, dict] = {
    "Title":     {"pt": 16, "zh_font": ZH_FONT_HEADING, "center": True},   # 三号
    "Heading 1": {"pt": 16, "zh_font": ZH_FONT_HEADING, "center": True},   # 文章标题：三号 居中
    "Heading 2": {"pt": 15, "zh_font": ZH_FONT_HEADING, "center": False},  # 章节标题：小三
    "Heading 3": {"pt": 15, "zh_font": ZH_FONT_HEADING, "center": False},
    "Heading 4": {"pt": 15, "zh_font": ZH_FONT_HEADING, "center": False},
    "Heading 5": {"pt": 15, "zh_font": ZH_FONT_HEADING, "center": False},
    "Heading 6": {"pt": 15, "zh_font": ZH_FONT_HEADING, "center": False},
}
BORDER_SIZE = "4"
BORDER_COLOR = "000000"


def set_run_font(run, size_pt: int, bold: bool | None = None,
                 zh_font: str = ZH_FONT_BODY, en_font: str = EN_FONT):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    rFonts.set(qn("w:eastAsia"), zh_font)
    rFonts.set(qn("w:cs"), en_font)

    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(size_pt * 2))
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size_pt * 2))


def paragraph_settings(paragraph) -> tuple[int, bool | None, str, bool]:
    """返回 (字号pt, 是否加粗, 中文字体, 是否居中)"""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in HEADING_CONFIG:
        cfg = HEADING_CONFIG[style_name]
        return cfg["pt"], True, cfg["zh_font"], cfg["center"]
    return BASE_PT, None, ZH_FONT_BODY, False


def style_paragraph(paragraph):
    size, bold, zh_font, center = paragraph_settings(paragraph)
    if not paragraph.runs:
        pPr = paragraph._p.get_or_add_pPr()
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)
    for run in paragraph.runs:
        set_run_font(run, size, bold, zh_font=zh_font)

    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in HEADING_CONFIG:
        # 标题段落启用"与下段同页"，避免末尾孤行或与下方图表被分页拆开
        paragraph.paragraph_format.keep_with_next = True
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), BORDER_SIZE)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), BORDER_SIZE)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)


def style_table(table):
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for p in cell.paragraphs:
                style_paragraph(p)
            for nested in cell.tables:
                style_table(nested)


_HEADING_STYLE_IDS = {"Title", "Heading1", "Heading2", "Heading3", "Heading4", "Heading5", "Heading6"}


def style_normal_definition(document):
    styles_element = document.styles.element
    for style in styles_element.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId")) or ""
        zh_font = ZH_FONT_HEADING if style_id in _HEADING_STYLE_IDS else ZH_FONT_BODY
        rPr = style.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), EN_FONT)
        rFonts.set(qn("w:hAnsi"), EN_FONT)
        rFonts.set(qn("w:eastAsia"), zh_font)
        rFonts.set(qn("w:cs"), EN_FONT)


def process_docx(path: Path):
    doc = Document(path)
    style_normal_definition(doc)
    for paragraph in doc.paragraphs:
        style_paragraph(paragraph)
    for table in doc.tables:
        style_table(table)
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is None:
                continue
            for p in header.paragraphs:
                style_paragraph(p)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is None:
                continue
            for p in footer.paragraphs:
                style_paragraph(p)
    doc.save(path)


def main():
    root = Path(sys.argv[1] if len(sys.argv) > 1 else "docs/招标/word")
    files = sorted(root.rglob("*.docx"))
    if not files:
        print(f"No docx files under {root}")
        return
    print(f"Found {len(files)} docx files under {root}")
    for f in files:
        print(f"  styling: {f.relative_to(root)}")
        try:
            process_docx(f)
        except Exception as e:
            print(f"    !! failed: {e}")
    print("All done.")


if __name__ == "__main__":
    main()
