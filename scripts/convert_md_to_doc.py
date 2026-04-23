#!/usr/bin/env python3
"""
将 Markdown 文档转换为 Word 或 PDF，并保持 Mermaid 图表可见
使用方法：
    python convert_md_to_doc.py <input.md> [--format word|pdf] [--output <output>] [--no-style]

默认转换完成后会自动调用 apply_docx_style.process_docx 对产物应用本项目
统一样式（宋体 14pt / 标题加粗 / 表格边框）。用 --no-style 可跳过。
"""

import re
import sys
import os
import base64
import urllib.parse
import urllib.request
import json
import tempfile
from pathlib import Path
from typing import List, Tuple, Optional

try:
    import markdown
    from markdown.extensions import tables, fenced_code
except ImportError:
    print("错误: 需要安装 markdown 库")
    print("运行: pip install markdown")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("运行: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
    import io
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    # 不在这里打印警告，因为可能不需要处理图片


def render_mermaid_to_image(mermaid_code: str, output_path: str, format: str = 'png') -> bool:
    """
    使用 Mermaid.ink API 或本地 mermaid-cli 将 Mermaid 代码渲染为图片
    """
    # 优先尝试使用本地 mermaid-cli（更可靠）
    try:
        import subprocess
        # 创建临时 mermaid 文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(mermaid_code)
            temp_mmd = f.name
        
        # 使用 mmdc 渲染（高分辨率：宽度 3200px，缩放 3x，背景透明，高质量）
        # -w: 宽度（像素）
        # -s: 缩放因子（2-4 之间，越高越清晰但文件越大）
        # -b: 背景色（transparent 透明）
        # -t: 主题（可选）
        # 尝试找到 mmdc 的完整路径
        mmdc_cmd = 'mmdc'
        npm_global_bin = os.path.expanduser('~/.npm-global/bin/mmdc')
        if os.path.exists(npm_global_bin):
            mmdc_cmd = npm_global_bin
        
        result = subprocess.run(
            [mmdc_cmd, '-i', temp_mmd, '-o', output_path, '-b', 'white', '-w', '3200', '-s', '3'],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        os.unlink(temp_mmd)
        
        if result.returncode == 0 and os.path.exists(output_path):
            print(f"  ✓ 使用本地 mmdc 渲染成功")
            return True
        else:
            print(f"  mmdc 返回码: {result.returncode}, stderr: {result.stderr[:200] if result.stderr else 'none'}")
    except FileNotFoundError:
        print(f"  mmdc 未找到，尝试在线服务...")
    except subprocess.TimeoutExpired:
        print(f"  警告: mermaid-cli 渲染超时")
    except Exception as e:
        print(f"  mmdc 异常: {e}")
    
    # 如果本地工具不可用，尝试在线服务
    try:
        import ssl
        # 创建不验证证书的上下文（仅用于开发环境）
        ssl_context = ssl.create_default_context()
        ssl_context.check_hostname = False
        ssl_context.verify_mode = ssl.CERT_NONE
        
        # 使用 mermaid.ink 在线服务（尝试获取更高分辨率的图片）
        # 注意：mermaid.ink 可能不支持自定义分辨率，但我们可以尝试
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        # 尝试使用 SVG 然后转换为 PNG（如果支持）
        url = f"https://mermaid.ink/img/{encoded}"
        
        # 下载图片
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, context=ssl_context, timeout=10) as response:
            with open(output_path, 'wb') as f:
                f.write(response.read())
        
        # 如果下载成功，尝试使用 PIL 提高分辨率（如果可用）
        if PIL_AVAILABLE and os.path.exists(output_path):
            try:
                img = Image.open(output_path)
                # 将图片放大 3 倍以提高清晰度，使用高质量重采样
                new_size = (img.width * 3, img.height * 3)
                img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
                # 保存为高质量 PNG，设置 300 DPI
                img_resized.save(output_path, 'PNG', quality=100, dpi=(300, 300))
                print(f"  📈 图片已放大: {img.size[0]}x{img.size[1]} → {new_size[0]}x{new_size[1]} (300 DPI)")
            except Exception as e:
                print(f"  ⚠️  图片处理失败: {e}")
                pass  # 如果处理失败，使用原始图片
        
        return True
    except Exception as e:
        print(f"  警告: 无法渲染 Mermaid 图表: {e}")
        print(f"  建议安装 mermaid-cli: npm install -g @mermaid-js/mermaid-cli")
        return False


def extract_mermaid_blocks(md_content: str) -> List[Tuple[str, str, int]]:
    """
    提取所有 Mermaid 代码块
    返回: [(mermaid_code, block_id, start_pos), ...]
    """
    pattern = r'```mermaid\n(.*?)\n```'
    matches = []
    
    for i, match in enumerate(re.finditer(pattern, md_content, re.DOTALL)):
        mermaid_code = match.group(1).strip()
        block_id = f"mermaid_{i+1}"
        matches.append((mermaid_code, block_id, match.start()))
    
    return matches


def replace_mermaid_with_images(md_content: str, image_dir: Path, format: str = 'png') -> Tuple[str, List[str]]:
    """
    将 Mermaid 代码块替换为图片引用
    返回: (修改后的 markdown, 图片路径列表)
    """
    mermaid_blocks = extract_mermaid_blocks(md_content)
    image_paths = []
    
    # 从后往前替换，避免位置偏移
    for mermaid_code, block_id, start_pos in reversed(mermaid_blocks):
        image_filename = f"{block_id}.{format}"
        image_path = image_dir / image_filename
        
        # 渲染 Mermaid 图表
        if render_mermaid_to_image(mermaid_code, str(image_path), format):
            # 替换为图片引用（使用相对于 image_dir 的路径，但只保留文件名）
            # 这样在 Word 中查找时更容易匹配
            image_md = f'\n![Mermaid 图表]({image_path.name})\n'
            # 找到对应的代码块结束位置
            pattern = r'```mermaid\n.*?\n```'
            match = re.search(pattern, md_content[start_pos:], re.DOTALL)
            if match:
                end_pos = start_pos + match.end()
                md_content = md_content[:start_pos] + image_md + md_content[end_pos:]
                image_paths.append(str(image_path))
                print(f"  ✅ 图表 {block_id} 渲染成功: {image_path.name}")
        else:
            # 如果渲染失败，保留原始代码块但添加注释
            print(f"  警告: 图表 {block_id} 渲染失败，保留原始代码")
    
    return md_content, image_paths


def add_paragraph_with_bold(doc: Document, text: str, style: Optional[str] = None):
    """
    将包含 **加粗** 标记的文本写入段落
    """
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
    return paragraph


def add_cell_text_with_bold(cell, text: str):
    """
    将包含 **加粗** 标记的文本写入表格单元格
    """
    cell.text = ''
    paragraph = cell.paragraphs[0]
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def apply_project_style(docx_path: str) -> bool:
    """
    调用 apply_docx_style.process_docx 为 docx 应用本项目统一样式。
    该模块与本脚本同目录，运行时 sys.path[0] 指向 scripts/ 可直接 import。
    """
    try:
        import apply_docx_style
    except ImportError as e:
        print(f"  ⚠️  无法导入 apply_docx_style（跳过样式处理）: {e}")
        return False

    try:
        apply_docx_style.process_docx(Path(docx_path))
        print(f"  🎨 已应用项目样式（宋体 14pt / 标题加粗 / 表格边框）")
        return True
    except Exception as e:
        print(f"  ⚠️  应用样式失败: {e}")
        return False


def markdown_to_word(md_file: str, output_file: str, image_dir: Path):
    """
    将 Markdown 转换为 Word 文档
    """
    print(f"读取 Markdown 文件: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 提取并渲染 Mermaid 图表
    print("处理 Mermaid 图表...")
    md_content, image_paths = replace_mermaid_with_images(md_content, image_dir)
    
    # 创建 Word 文档
    doc = Document()

    # 收窄页边距，给大图（尤其是 mermaid 长图）留出更多可用空间。
    # A4 为 8.27"×11.69"，边距 0.5" 后可用区 ~7.27"×10.69"。
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # 解析 Markdown
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('# ').strip()
            if text:
                heading = doc.add_heading(text, level=min(level, 9))
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 图片
        elif line.startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if match:
                alt_text, img_path = match.groups()
                # 尝试多个路径查找图片文件
                img_full_path = None
                
                # 路径1: 相对于 markdown 文件所在目录
                test_path = Path(md_file).parent / img_path
                if test_path.exists():
                    img_full_path = test_path
                else:
                    # 路径2: 相对于 image_dir（绝对路径）
                    test_path = image_dir / Path(img_path).name
                    if test_path.exists():
                        img_full_path = test_path
                    else:
                        # 路径3: 直接使用 img_path（如果是绝对路径）
                        test_path = Path(img_path)
                        if test_path.exists():
                            img_full_path = test_path
                        else:
                            # 路径4: 在 image_dir 中查找同名文件
                            img_name = Path(img_path).name
                            test_path = image_dir / img_name
                            if test_path.exists():
                                img_full_path = test_path
                
                if img_full_path and img_full_path.exists():
                    try:
                        # 大图允许占用页边距区域。配合上面设置的 0.5" 边距，
                        # A4 可用区 ~7.27"×10.69"；宽度放到 7.5"（略压边距）。
                        # 高度上限收到 8.5"，给"标题 + 1-3 行导语 + 段距 + 页眉缓冲"
                        # 留约 2.2" 空间，确保 Word/WPS 各自默认页眉页脚距离下
                        # 也能稳定同页，不触发 keep_with_next 失效的边界情况。
                        MAX_W_IN = 7.5
                        MAX_H_IN = 8.5
                        display_width = None
                        display_height = None
                        img_w_px = img_h_px = None

                        if PIL_AVAILABLE:
                            try:
                                img = Image.open(img_full_path)
                                img_w_px, img_h_px = img.size
                                aspect = img_h_px / img_w_px  # h / w
                                # 先按最大宽度铺，超出最大高度则改按高度约束
                                w_in = MAX_W_IN
                                h_in = w_in * aspect
                                if h_in > MAX_H_IN:
                                    h_in = MAX_H_IN
                                    w_in = h_in / aspect
                                display_width = Inches(w_in)
                                display_height = Inches(h_in)
                            except Exception:
                                display_width = Inches(MAX_W_IN)

                        if display_width is None:
                            display_width = Inches(MAX_W_IN)

                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 压缩图片段上下段距，腾出纵向空间给标题共页
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        run = paragraph.add_run()
                        if display_height is not None:
                            picture = run.add_picture(str(img_full_path), width=display_width, height=display_height)
                        else:
                            picture = run.add_picture(str(img_full_path), width=display_width)

                        # 让图片段自身尽量不跨页；并把紧挨其上的段（标题 + 短导语）
                        # 标记为与下段同页。向上回溯最多 3 段，遇到标题则包含该段后停止，
                        # 避免把无关内容串连过去。规则目标："标题 + 不超过 3 行正文 + 图"同页。
                        paragraph.paragraph_format.keep_together = True
                        MAX_LOOKBACK = 3
                        total = len(doc.paragraphs)
                        back = 2  # paragraphs[-1] 是图片段本身
                        steps = 0
                        while back <= total and steps < MAX_LOOKBACK:
                            prev = doc.paragraphs[-back]
                            prev.paragraph_format.keep_with_next = True
                            prev.paragraph_format.space_after = Pt(0)
                            style_name = prev.style.name if prev.style else ""
                            steps += 1
                            # 遇到标题（或 Title）就停：已经把整个"小节头"粘到图上
                            if style_name.startswith("Heading") or style_name == "Title":
                                break
                            back += 1

                        if img_w_px is not None:
                            print(f"  ✅ 插入图片: {img_full_path.name} (源 {img_w_px}x{img_h_px}px → 显示 {display_width.inches:.2f}x{display_height.inches:.2f}英寸)")
                        else:
                            print(f"  ✅ 插入图片: {img_full_path.name} (显示宽度 {display_width.inches:.2f}英寸)")
                    except Exception as e:
                        print(f"  ⚠️  无法插入图片 {img_path}: {e}")
                        doc.add_paragraph(f"[图片: {alt_text}]")
                else:
                    print(f"  ⚠️  图片未找到: {img_path}")
                    doc.add_paragraph(f"[图片未找到: {img_path}]")
        
        # 表格
        elif '|' in line and line.count('|') >= 2:
            # 收集表格行
            table_lines = [line]
            i += 1
            # 跳过分隔行
            if i < len(lines) and '|' in lines[i] and '---' in lines[i]:
                i += 1
            # 收集数据行
            while i < len(lines) and '|' in lines[i] and lines[i].count('|') >= 2:
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # 回退一行
            
            # 创建表格
            if len(table_lines) > 0:
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # 添加表头
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        add_cell_text_with_bold(header_cells[j], header)
                    
                    # 添加数据行
                    for row_line in table_lines[1:]:
                        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        if len(cells) == len(headers):
                            row = table.add_row()
                            for j, cell in enumerate(cells):
                                add_cell_text_with_bold(row.cells[j], cell)
        
        # 引用块
        elif line.startswith('>'):
            text = line.lstrip('> ').strip()
            if text:
                add_paragraph_with_bold(doc, text, style='Quote')
        
        # 分隔线
        elif line.startswith('---') or line.startswith('***'):
            # 分隔线不转，避免 Word 中过长换行
            pass
        
        # 普通段落
        elif line:
            add_paragraph_with_bold(doc, line)
        
        i += 1
    
    # 保存文档
    print(f"保存 Word 文档: {output_file}")
    doc.save(output_file)
    print(f"✅ Word 转换完成！")
    print(f"   输出文件: {output_file}")
    if image_paths:
        print(f"   生成的图片: {len(image_paths)} 个")


def convert_word_to_pdf(word_file: str, pdf_file: str):
    """
    将 Word 文档转换为 PDF
    """
    # 方法1: 尝试使用 docx2pdf（macOS 需要 Microsoft Word）
    try:
        import docx2pdf
        print("使用 docx2pdf 转换（需要 Microsoft Word）...")
        docx2pdf.convert(word_file, pdf_file)
        if os.path.exists(pdf_file):
            print(f"✅ PDF 转换完成: {pdf_file}")
            return
    except ImportError:
        print("docx2pdf 未安装，尝试其他方法...")
    except Exception as e:
        print(f"docx2pdf 转换失败: {e}")
    
    # 方法1.5: 尝试使用 AppleScript 直接调用 Word（macOS）
    try:
        import subprocess
        word_path = os.path.abspath(word_file)
        pdf_path = os.path.abspath(pdf_file)
        
        # 使用正确的 AppleScript 语法
        script = f'''
tell application "Microsoft Word"
    set wordFile to POSIX file "{word_path}" as string
    open wordFile
    set theDoc to active document
    set pdfFile to POSIX file "{pdf_path}" as string
    save as theDoc file name pdfFile file format format PDF
    close theDoc saving no
end tell
'''
        result = subprocess.run(
            ['osascript', '-e', script],
            capture_output=True,
            text=True,
            timeout=60
        )
        if os.path.exists(pdf_file):
            print(f"✅ PDF 转换完成 (使用 AppleScript + Word): {pdf_file}")
            return
        elif result.stderr:
            print(f"AppleScript 转换失败: {result.stderr}")
    except Exception as e:
        print(f"AppleScript 转换失败: {e}")
    
    # 方法2: 尝试使用 pandoc
    try:
        import subprocess
        result = subprocess.run(
            ['pandoc', word_file, '-o', pdf_file, '--pdf-engine=xelatex'],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0 and os.path.exists(pdf_file):
            print(f"✅ PDF 转换完成 (使用 pandoc): {pdf_file}")
            return
        else:
            print(f"pandoc 转换失败: {result.stderr}")
    except FileNotFoundError:
        print("pandoc 未安装")
    except Exception as e:
        print(f"pandoc 转换失败: {e}")
    
    # 方法3: 提示用户手动转换
    print("\n⚠️  自动 PDF 转换失败")
    print("请使用以下方法之一手动转换:")
    print(f"  1. 打开 Word 文档: {word_file}")
    print("     文件 → 另存为 → 选择 PDF 格式")
    print(f"  2. 或安装工具后重试:")
    print("     - macOS: pip3 install docx2pdf")
    print("     - 或: brew install pandoc")
    print(f"\nWord 文档已保存: {word_file}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    input_file = sys.argv[1]
    if not os.path.exists(input_file):
        print(f"错误: 文件不存在: {input_file}")
        sys.exit(1)
    
    # 解析参数
    output_format = 'word'
    output_file = None
    apply_style = True

    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '--format':
            i += 1
            if i < len(sys.argv):
                output_format = sys.argv[i]
        elif sys.argv[i] == '--output' or sys.argv[i] == '-o':
            i += 1
            if i < len(sys.argv):
                output_file = sys.argv[i]
        elif sys.argv[i] == '--no-style':
            apply_style = False
        i += 1
    
    # 确定输出文件名
    if not output_file:
        base_name = Path(input_file).stem
        if output_format == 'word':
            output_file = f"{base_name}.docx"
        elif output_format == 'pdf':
            output_file = f"{base_name}.pdf"
        else:
            output_file = f"{base_name}.docx"
    
    # 创建图片目录（在输出文件同目录）
    output_path = Path(output_file)
    image_dir = output_path.parent / "mermaid_images"
    image_dir.mkdir(parents=True, exist_ok=True)
    
    if output_format == 'word':
        markdown_to_word(input_file, output_file, image_dir)
        if apply_style:
            apply_project_style(output_file)
    elif output_format == 'pdf':
        # 先转换为 Word，再转换为 PDF
        word_file = output_file.replace('.pdf', '.docx')
        print("第一步: 转换为 Word 文档...")
        markdown_to_word(input_file, word_file, image_dir)
        if apply_style:
            apply_project_style(word_file)

        print("\n第二步: 转换为 PDF 文档...")
        convert_word_to_pdf(word_file, output_file)
    else:
        print(f"错误: 不支持的格式: {output_format}")
        print("支持的格式: word, pdf")


if __name__ == '__main__':
    main()
